#!/usr/bin/env python3
"""
api_util/digital_to_json.py — digital-born PDF/DOCX → `atrium_document` JSON.

The `digital-convert` originator (hub issue #18 §1a, plan §2). `BLOCK_OWNERS` gives the
positional plane — `pages`, `content`, `lines`, `tables` — two possible originators, and
which one applies is fixed per document by `source.origin`:

    ocr:* / vlm:* / ABBYY-ALTO   → alto-postprocess
    digital-born* / pdf / docx   → digital-convert   ← this module

The two are mutually exclusive per document, so nothing here is a second opinion on an
ALTO document; it is the *only* writer of the plane for a born-digital one.

Four layers, per `agent_dev_logs/plans/18.plan.md` §2:

  A. **Extraction** — `pdfplumber` for PDF geometry, `python-docx` for OOXML, into the
     internal `DigitalDocument`. The only layer that needs `requirements_digital.txt`;
     imported lazily so the base install and the fast lane never pay for it.
  B. **Normalization** — paragraph grouping from geometry, and the decode-sanity gate
     that turns "born-digital" into "born-digital and actually readable".
  C. **Serialization** — into `DocumentRecord` as `program="digital-convert"`, writing
     only this program's declared field grants.
  D. **Validation & export** — the §1b round-trip assertion *and*
     `validate_document()`. If either fails, **no `doc.json` is emitted**.

## Born-digital ≠ trustworthy text

The reason layer B exists at all. Issue #10's research pass found PDFs whose text layer
extracts *successfully* and *plausibly*, and is wrong: non-embedded WinAnsi Helvetica with
no `/ToUnicode` map, so CP1250 (Central European) bytes are read as CP1252 (Western).

    intended : Zpráva o sondě číslo 3. Nalezeny hřeby, vrstva ornice měla
    extracted: Zpráva o sondì èíslo 3. Nalezeny høeby, vrstva ornice mìla

No exception, no encoding error, no length change — and `Zpráva` survives intact, because
`á` sits at the same codepoint in both encodings. *Partially* correct output is what makes
this dangerous: a downstream reader has no signal that anything went wrong.

`CP1250_MISREADS` is derived from the two codecs at import time rather than typed out by
hand, so it is exactly the set of byte values where the two disagree *and* the CP1250
reading is a Czech letter — provably complete, and self-documenting about why `á`/`é` are
absent from it.

Detection feeds two outputs, and deliberately not a third:

  * `lines[].categ` = **`"Garbage"`** — the exact spelling `api_util/json_to_md.py`'s
    `DROP_CATEGORIES` filters on. `categ` is an open string in the schema, so a synonym
    would silently disable the filter instead of failing validation.
  * `pages[].needs_ocr` + `pages[].needs_ocr_reason` — the documented digital→OCR
    hand-off. Setting `needs_ocr` is what *authorises* `alto-postprocess` to re-originate
    this record's plane, even though `source.origin` stays `digital-born-*` (which remains
    truthful: it describes how the ORIGINAL was acquired). `pages[].ocr` is never granted
    to this program, so "was this OCR'd" stays answerable from the record.
  * **Not** a silent repair. `decode_sanity()` returns the recovered string because it is
    the single most useful thing a human triaging the document can see, but this converter
    REPORTS; routing policy lives outside it. Rewriting text under the reader's feet would
    make `source.sha256` describe a document whose text no longer matches it.

## Coordinates

`$defs/bbox` is normative: **origin top-left, y increasing downwards**, in the unit named by
that page's `canvas.unit`. For pdfplumber that means `top`/`bottom` — never `y0`/`y1`, which
are PDF user space (origin bottom-left) and are the specific mistake the schema description
calls out. DOCX has no page geometry without rendering, so a DOCX document carries no
`bbox` and therefore no `canvas` either.

Usage:
    python api_util/digital_to_json.py report.pdf --out report.document.json
    python api_util/digital_to_json.py report.docx --document-json prev.json --out out.json
"""

from __future__ import annotations

import argparse
import hashlib
import os
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Tuple

# The vendored hub canonical modules live at the repo root; this file is in api_util/.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from atrium_document import (  # noqa: E402
    DocumentRecord,
    canonical_doc_id,
    validate_document,
)

PROGRAM = "digital-convert"

#: `source.origin` values this converter writes. Both match an `ORIGIN_ORIGINATORS` prefix,
#: which is what routes the positional plane to this program. An origin the table has not
#: been taught makes `_assert_origin_consistent()` ABSTAIN rather than refuse — silently
#: switching §1a off for that document — so these strings are not cosmetic.
ORIGIN_PDF = "digital-born-pdf"
ORIGIN_DOCX = "digital-born-docx"

#: The two load-bearing `categ` spellings (`json_to_md.DROP_CATEGORIES`).
CATEG_GARBAGE = "Garbage"
CATEG_INVERTED = "Inverted"

#: Below this decode-sanity ratio a line is `Garbage` — the density signal, for a line
#: whose diacritics are mostly wrong.
QUALITY_GARBAGE_BELOW = 0.90

#: ...and this many confusable characters condemn a line regardless of its length.
#:
#: A ratio ALONE provably cannot catch the real corruption. The observed line
#: "Zpráva o sondì èíslo 3. Nalezeny høeby, vrstva ornice mìla" carries 4 misreads across
#: 46 letters, which scores 0.913 — comfortably above any ratio cut loose enough not to
#: condemn clean text. Czech simply does not contain `ì`, `è`, `ø` or `ù`, so their presence
#: is not a density question: ONE may be a foreign word in a quotation, but TWO in a single
#: line is a systematic decode fault, which is exactly the shape of this failure.
#:
#: The two rules are kept separate rather than folded into one score because they answer
#: different questions — "how corrupt is this line?" (which `quality_score` must keep
#: reporting as a 0–1 axis) and "is this line trustworthy at all?".
GARBAGE_MIN_HITS = 2

#: Page-level band cutoffs, matching the vocabulary alto-postprocess already uses so a
#: consumer does not have to know which originator wrote the page.
BAND_CLEAR_AT = 0.90
BAND_NOISY_AT = 0.50

#: Paragraph splitting. Measured on tests/fixtures/digital/minimal.pdf: 2.0 pt between
#: lines of one paragraph, 34.0 pt between paragraphs (digital_born/README.md). A ratio
#: is used rather than an absolute, because the gap scales with the font.
PARAGRAPH_GAP_RATIO = 1.8


def _czech_letters() -> frozenset:
    return frozenset("áčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ")


def _build_cp1250_misreads() -> Dict[str, str]:
    """Characters a CP1250 byte turns into when read as CP1252, where CP1250 meant Czech.

    Derived from the codecs, not hand-typed: for every high byte, decode it both ways and
    keep the pair only when the readings differ AND the CP1250 reading is a Czech letter.
    That makes the table provably complete for this failure mode, and explains its own
    omissions — `á` (0xE1) and `é` (0xE9) decode identically under both codecs, which is
    exactly why "Zpráva" survives a mis-decode while "sondě" does not.
    """
    czech = _czech_letters()
    table: Dict[str, str] = {}
    for byte in range(0x80, 0x100):
        try:
            western = bytes([byte]).decode("cp1252")
            eastern = bytes([byte]).decode("cp1250")
        except UnicodeDecodeError:  # undefined in one of the codecs
            continue
        if western != eastern and eastern in czech:
            table[western] = eastern
    return table


CP1250_MISREADS: Dict[str, str] = _build_cp1250_misreads()


# ──────────────────────────────────────────────────────────────────────────────
# Internal representation (shared by both Layer A extractors)
# ──────────────────────────────────────────────────────────────────────────────


@dataclass
class DigitalLine:
    """One text line, before it becomes a `lines[]` row."""

    page: str
    line: int
    text: str
    bbox: Optional[List[float]] = None
    font: str = ""
    size: float = 0.0
    bold: bool = False
    italic: bool = False
    heading_level: Optional[int] = None
    inverted: bool = False
    group_id: Optional[str] = None
    categ: Optional[str] = None
    quality_score: Optional[float] = None
    lang: Optional[str] = None


@dataclass
class DigitalTable:
    """One table grid. `cells[].group_id` is the join key back into `lines[]`."""

    table_id: str
    page: str
    caption: str = ""
    n_rows: int = 0
    n_cols: int = 0
    group_id: Optional[str] = None
    cells: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class DigitalPage:
    page: str
    page_index: int
    width: float = 0.0
    height: float = 0.0
    unit: str = "pt"
    lines: List[DigitalLine] = field(default_factory=list)
    tables: List[DigitalTable] = field(default_factory=list)
    needs_ocr: bool = False
    needs_ocr_reason: str = ""
    quality_score: Optional[float] = None
    quality_band: Optional[str] = None

    @property
    def has_geometry(self) -> bool:
        return any(line.bbox for line in self.lines)


@dataclass
class DigitalDocument:
    doc_id: str
    origin: str
    media_type: str
    pages: List[DigitalPage] = field(default_factory=list)
    sha256: str = ""
    filename: str = ""
    reading_order: str = "layout"

    def all_lines(self) -> List[DigitalLine]:
        return [line for page in self.pages for line in page.lines]


# ──────────────────────────────────────────────────────────────────────────────
# Layer B — normalization and the decode-sanity gate  (pure; no optional deps)
# ──────────────────────────────────────────────────────────────────────────────


@dataclass
class DecodeReport:
    """What `decode_sanity()` concluded about one string."""

    score: float
    suspicious: int
    letters: int
    recovered: Optional[str] = None

    @property
    def is_garbage(self) -> bool:
        """Either signal condemns the line — see `GARBAGE_MIN_HITS` for why both are needed."""
        return self.suspicious >= GARBAGE_MIN_HITS or self.score < QUALITY_GARBAGE_BELOW


def decode_sanity(text: str) -> DecodeReport:
    """Score 0..1 for "this text decoded correctly", plus the recovered reading.

    The score is the share of letters that are NOT a CP1250→CP1252 misread. It is
    deliberately a ratio, not a count: a caption of three words and a page of three hundred
    must be comparable, and `pages[].quality_score` is documented as a 0–1 axis either way.

    `recovered` is offered only when the whole string round-trips — `encode("cp1252")` then
    `decode("cp1250")`. A string containing a character CP1252 cannot represent was never a
    CP1250 byte sequence read this way, so a partial guess there would be fabrication.
    """
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return DecodeReport(score=1.0, suspicious=0, letters=0)

    suspicious = sum(1 for c in letters if c in CP1250_MISREADS)
    score = 1.0 - (suspicious / len(letters))

    recovered: Optional[str] = None
    if suspicious:
        try:
            candidate = text.encode("cp1252").decode("cp1250")
        except (UnicodeEncodeError, UnicodeDecodeError):
            candidate = None
        if candidate and candidate != text:
            recovered = candidate

    return DecodeReport(score=score, suspicious=suspicious, letters=len(letters), recovered=recovered)


def classify_line(line: DigitalLine, report: DecodeReport) -> Optional[str]:
    """The `categ` this line carries, or None to leave the field absent.

    Only the two spellings `json_to_md.DROP_CATEGORIES` knows are ever produced. `Inverted`
    wins over `Garbage`: mirrored text is a rendering-level fault, and its extracted string
    is not evidence of anything, so reporting a decode verdict on it would be noise.
    """
    if line.inverted:
        return CATEG_INVERTED
    if report.is_garbage:
        return CATEG_GARBAGE
    return None


def assign_group_ids(lines: Sequence[DigitalLine]) -> None:
    """Group consecutive lines into paragraphs, in place, from vertical gaps.

    `lines[].group_id` is what `json_to_md` turns back into a blank line, so this is the
    only thing preserving paragraph structure through the JSON. Lines without geometry
    (DOCX) are left ungrouped — python-docx already gives one paragraph per line, and
    inventing groups from nothing would be worse than the absence the schema allows.

    The threshold is relative (`PARAGRAPH_GAP_RATIO` × the median gap) so it holds for a
    10 pt body and a 24 pt heading alike; an absolute cut tuned on one fixture would not.
    """
    positioned = [ln for ln in lines if ln.bbox]
    if len(positioned) < 2:
        return

    gaps: List[float] = []
    for previous, current in zip(positioned, positioned[1:], strict=False):
        gaps.append(max(0.0, current.bbox[1] - previous.bbox[3]))

    # The baseline is a LOW PERCENTILE of the positive gaps, not their median: within-
    # paragraph leading is the common small value, and paragraph breaks are the rare large
    # one, so the median is dragged upward by them exactly when there are few lines. On the
    # measured fixture (gaps 2.0 and 34.0) the median is 34.0 and nothing ever splits — the
    # threshold would sit above the very break it exists to find. The 25th percentile also
    # resists a single freak sub-point gap, which a bare `min()` would not.
    ordered = sorted(g for g in gaps if g > 0)
    baseline = ordered[max(0, len(ordered) // 4)] if ordered else 0.0
    threshold = baseline * PARAGRAPH_GAP_RATIO if baseline else None

    group_index = 0
    positioned[0].group_id = f"p{group_index}"
    for gap, current in zip(gaps, positioned[1:], strict=True):  # one gap per following line
        if threshold is not None and gap > threshold:
            group_index += 1
        current.group_id = f"p{group_index}"


def assess_page(page: DigitalPage) -> None:
    """Fill in `quality_score`, `quality_band`, `needs_ocr` and its reason, in place.

    `needs_ocr_reason` is written whenever `needs_ocr` is, and never alone. The two mean
    OPPOSITE things on the two originator paths — "no text layer" for a scan, "a text layer
    that lies" here — and `json_to_md` emits the reason as a cue the model reads. Without
    it every digital-born page rendered "no extractable text layer", which is false by
    definition for a document that has one.
    """
    scored = [ln.quality_score for ln in page.lines if ln.quality_score is not None]
    page.quality_score = round(sum(scored) / len(scored), 4) if scored else 1.0

    if page.quality_score >= BAND_CLEAR_AT:
        page.quality_band = "Clear"
    elif page.quality_score >= BAND_NOISY_AT:
        page.quality_band = "Noisy"
    else:
        page.quality_band = "Trash"

    garbled = [ln for ln in page.lines if ln.categ == CATEG_GARBAGE]
    if garbled:
        page.needs_ocr = True
        page.needs_ocr_reason = (
            f"embedded text layer does not decode: {len(garbled)} of {len(page.lines)} lines "
            f"carry CP1250 bytes read as CP1252 (mojibake diacritics), page decode-sanity "
            f"{page.quality_score:.2f}. The page has a text layer; it is not trustworthy."
        )


def normalize(doc: DigitalDocument) -> DigitalDocument:
    """Run Layer B over an extracted document, in place, and return it."""
    for page in doc.pages:
        for line in page.lines:
            report = decode_sanity(line.text)
            line.quality_score = round(report.score, 4)
            categ = classify_line(line, report)
            if categ:
                line.categ = categ
        assign_group_ids(page.lines)
        assess_page(page)
    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Layer A — extraction (lazy optional imports)
# ──────────────────────────────────────────────────────────────────────────────


def _sha256(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1 << 20), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _missing(dependency: str, module: str) -> RuntimeError:
    return RuntimeError(
        f"{dependency} is required to convert this format but is not installed. "
        f"Install the converter stack: pip install -r requirements_digital.txt "
        f"(missing import: {module})"
    )


def _font_flags(fontname: str) -> Tuple[bool, bool]:
    lowered = (fontname or "").lower()
    return ("bold" in lowered), ("italic" in lowered or "oblique" in lowered)


def _is_inverted(matrix: Any) -> bool:
    """True when a glyph's transformation matrix mirrors or flips it.

    pdfplumber exposes the text matrix as `(a, b, c, d, e, f)`. A negative `a` mirrors
    horizontally and a negative `d` flips vertically; either makes the extracted string
    unreliable as *text* however cleanly it decodes, which is what `Inverted` records.
    """
    if not matrix or len(matrix) < 4:
        return False
    try:
        return float(matrix[0]) < 0 or float(matrix[3]) < 0
    except (TypeError, ValueError):
        return False


def extract_pdf(path: str, doc_id: Optional[str] = None) -> DigitalDocument:
    """Layer A for PDF, via pdfplumber (MIT).

    Chars are grouped into lines by their `top` coordinate rather than trusting a
    line-extraction helper, so the bbox this writes is built from the same numbers the
    grouping used. `top`/`bottom` are pdfplumber's top-left-origin pair — the convention
    `$defs/bbox` requires — and `y0`/`y1` are deliberately never read.
    """
    try:
        import pdfplumber  # noqa: PLC0415  (optional dependency, imported on use)
    except ImportError as exc:
        raise _missing("pdfplumber", "pdfplumber") from exc

    document = DigitalDocument(
        doc_id=doc_id or canonical_doc_id(path),
        origin=ORIGIN_PDF,
        media_type="application/pdf",
        sha256=_sha256(path),
        filename=os.path.basename(path),
    )

    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages):
            label = str(page.page_number if page.page_number is not None else index + 1)
            current = DigitalPage(
                page=label,
                # 1-BASED: the schema types page_index as `minimum: 1` and documents it as
                # the "1-based physical position". A 0-based value is not merely unusual,
                # it fails validation — caught by the Layer D gate on the first run.
                page_index=index + 1,
                width=float(page.width or 0.0),
                height=float(page.height or 0.0),
                unit="pt",
            )

            rows: Dict[int, List[Any]] = {}
            for char in page.chars:
                rows.setdefault(round(float(char["top"])), []).append(char)

            for line_no, top in enumerate(sorted(rows)):
                chars = sorted(rows[top], key=lambda c: float(c["x0"]))
                text = "".join(c.get("text", "") for c in chars).strip()
                if not text:
                    continue
                fontname = str(chars[0].get("fontname", ""))
                bold, italic = _font_flags(fontname)
                current.lines.append(
                    DigitalLine(
                        page=label,
                        line=line_no,
                        text=text,
                        bbox=[
                            round(min(float(c["x0"]) for c in chars), 3),
                            round(min(float(c["top"]) for c in chars), 3),
                            round(max(float(c["x1"]) for c in chars), 3),
                            round(max(float(c["bottom"]) for c in chars), 3),
                        ],
                        font=fontname,
                        size=round(float(chars[0].get("size", 0.0)), 3),
                        bold=bold,
                        italic=italic,
                        inverted=any(_is_inverted(c.get("matrix")) for c in chars),
                    )
                )
            document.pages.append(current)

    return document


def extract_docx(path: str, doc_id: Optional[str] = None) -> DigitalDocument:
    """Layer A for DOCX, via python-docx (MIT).

    A DOCX has no page geometry until something renders it, so this writes NO `bbox` and
    therefore no `canvas` — the schema requires `canvas.unit` only when a bbox is present,
    and inventing coordinates would be worse than the honest absence. Everything lands on
    page "1" for the same reason: the file does not say where the page breaks fall.
    """
    try:
        import docx  # noqa: PLC0415  (optional dependency, imported on use)
    except ImportError as exc:
        raise _missing("python-docx", "docx") from exc

    document = DigitalDocument(
        doc_id=doc_id or canonical_doc_id(path),
        origin=ORIGIN_DOCX,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        sha256=_sha256(path),
        filename=os.path.basename(path),
        reading_order="flow",
    )
    page = DigitalPage(page="1", page_index=1, unit="pt")  # 1-based, per the schema
    source = docx.Document(path)

    for line_no, paragraph in enumerate(source.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style = (getattr(paragraph.style, "name", "") or "").strip()
        heading_level = None
        if style.lower().startswith("heading"):
            tail = style.split()[-1]
            heading_level = int(tail) if tail.isdigit() else 1
        runs = list(paragraph.runs)
        page.lines.append(
            DigitalLine(
                page="1",
                line=line_no,
                text=text,
                bold=bool(runs) and all(bool(r.bold) for r in runs),
                italic=bool(runs) and all(bool(r.italic) for r in runs),
                heading_level=heading_level,
            )
        )

    for table_no, table in enumerate(source.tables):
        grid = DigitalTable(
            table_id=f"t{table_no}",
            page="1",
            n_rows=len(table.rows),
            n_cols=len(table.columns),
            group_id=f"tbl{table_no}",
        )
        for row_no, row in enumerate(table.rows):
            for col_no, cell in enumerate(row.cells):
                grid.cells.append(
                    {
                        "row": row_no,
                        "col": col_no,
                        "is_header": row_no == 0,
                        "group_id": f"tbl{table_no}-r{row_no}c{col_no}",
                        "text": (cell.text or "").strip(),
                    }
                )
        page.tables.append(grid)

    document.pages.append(page)
    return document


def extract(path: str, doc_id: Optional[str] = None) -> DigitalDocument:
    """Dispatch on the file extension. Unknown formats fail loudly rather than guessing."""
    suffix = os.path.splitext(path)[1].lower()
    if suffix == ".pdf":
        return extract_pdf(path, doc_id=doc_id)
    if suffix in (".docx", ".docm"):
        return extract_docx(path, doc_id=doc_id)
    raise ValueError(f"unsupported input {path!r}: expected .pdf or .docx")


# ──────────────────────────────────────────────────────────────────────────────
# Layer C — serialization into DocumentRecord
# ──────────────────────────────────────────────────────────────────────────────


def _page_rows(doc: DigitalDocument) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for page in doc.pages:
        row: Dict[str, Any] = {
            "page": page.page,
            "page_index": page.page_index,
            "quality_score": page.quality_score,
            "quality_band": page.quality_band,
        }
        # `canvas.unit` is mandatory whenever any bbox exists on the page, and meaningless
        # when none does (DOCX) — so the whole block is conditional on real geometry.
        if page.has_geometry and page.width and page.height:
            row["canvas"] = {"width": page.width, "height": page.height, "unit": page.unit}
        if page.needs_ocr:
            row["needs_ocr"] = True
            row["needs_ocr_reason"] = page.needs_ocr_reason
        rows.append(row)
    return rows


def _line_rows(doc: DigitalDocument) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for line in doc.all_lines():
        row: Dict[str, Any] = {"page": line.page, "line": line.line, "text": line.text}
        if line.bbox:
            row["bbox"] = line.bbox
        if line.group_id is not None:
            row["group_id"] = line.group_id
        if line.categ:
            row["categ"] = line.categ
        if line.quality_score is not None:
            row["quality_score"] = line.quality_score
        if line.lang:
            row["lang"] = line.lang
        # Semantic style only — bold/italic/heading_level. Typeface and point size are
        # deliberately dropped: a reader can act on "this was a heading" and cannot act on
        # "this was Helvetica 12pt", and heading-ness partly duplicates `categ` already.
        style = {}
        if line.bold:
            style["bold"] = True
        if line.italic:
            style["italic"] = True
        if line.heading_level:
            style["heading_level"] = line.heading_level
        if style:
            row["style"] = style
        rows.append(row)
    return rows


def _table_rows(doc: DigitalDocument) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for page in doc.pages:
        for table in page.tables:
            rows.append(
                {
                    "table_id": table.table_id,
                    "page": table.page,
                    "caption": table.caption,
                    "n_rows": table.n_rows,
                    "n_cols": table.n_cols,
                    "group_id": table.group_id,
                    "cells": table.cells,
                }
            )
    return rows


def to_record(
    doc: DigitalDocument,
    baseline: Optional[str] = None,
    run_id: Optional[str] = None,
    paradata_ref: str = "",
    out_dir: str = ".",
    strict: bool = False,
) -> Tuple[DocumentRecord, List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Build the record. Returns it with the page and line rows, for Layer D to re-check.

    `DocumentRecord.open()` handles the baseline for us: a missing or absent path is not an
    error (rule 3, standalone-safe), and when one IS given every other tool's blocks are
    deep-copied through untouched (rule 2). Note that `open()` also decides the doc_id — a
    baseline's own id wins over the one derived here, because our id came from a LOCAL
    FILENAME and the record already knows what document it is.

    `set_source()` is called FIRST. The module re-checks blocks written before an origin
    arrives, so this is no longer strictly required — but `source.origin` is what authorises
    this program to write the plane at all, and writing it first keeps the authorisation
    visible in the code rather than dependent on a deferred re-check. It is also
    first-writer-wins, so on a re-run the original acquisition facts are preserved and only
    a genuine contradiction is reported.
    """
    record = DocumentRecord.open(
        doc.doc_id,
        PROGRAM,
        baseline=baseline,
        run_id=run_id,
        paradata_ref=paradata_ref,
        out_dir=out_dir,
        strict=strict,
    )
    record.set_source(
        sha256=doc.sha256,
        filename=doc.filename,
        media_type=doc.media_type,
        origin=doc.origin,
        page_count=len(doc.pages),
    )

    page_rows = _page_rows(doc)
    line_rows = _line_rows(doc)

    # pages[] and lines[] are FIELD-SPLIT (page-classification contributes category;
    # nlp-enrich contributes morphology), so merge_block — set_block would erase a
    # co-contributor's fields on a re-run over an existing record.
    record.merge_block("pages", page_rows)
    record.merge_block("lines", line_rows)

    # `content` has a single owner per document and `tables` is declared for the two
    # ORIGINATORS only — which are mutually exclusive per record — so neither has a
    # co-contributor to erase and set_block is both correct and quiet for them.
    body = "\n".join(line.text for line in doc.all_lines() if line.categ != CATEG_GARBAGE)
    record.set_block("content", {"text": body or None, "reading_order": doc.reading_order})

    tables = _table_rows(doc)
    if tables:
        record.set_block("tables", tables)

    return record, page_rows, line_rows


# ──────────────────────────────────────────────────────────────────────────────
# Layer D — validation and export
# ──────────────────────────────────────────────────────────────────────────────


def emit(
    record: DocumentRecord,
    page_rows: List[Dict[str, Any]],
    line_rows: List[Dict[str, Any]],
    out_path: Optional[str] = None,
) -> str:
    """The output gate: round-trip assertion, then schema validation, then write.

    Both checks run BEFORE `finalize()`, so a record that fails either is never written —
    not written-then-flagged. They catch different things and neither subsumes the other:

      * `assert_fields_survived()` catches a **field-ownership drop**, which the schema
        structurally cannot. `lines[]` requires only `page`+`line`, so a row stripped of its
        `text` by a too-narrow grant is a *valid* row. That exact bug shipped once already:
        an earlier draft granted this program only `["group_id"]` on `lines`, `merge_block()`
        honoured it silently, and the resulting text-free records validated clean.
      * `validate_document()` catches everything the schema does describe — and it raises
        rather than passing when `jsonschema` is absent, because a gate that quietly
        no-ops is indistinguishable from a passing one.
    """
    record.assert_fields_survived("lines", line_rows)
    record.assert_fields_survived("pages", page_rows)
    validate_document(record.to_dict())
    return record.finalize(out_path)


def convert(
    input_path: str,
    out_path: Optional[str] = None,
    baseline: Optional[str] = None,
    doc_id: Optional[str] = None,
    run_id: Optional[str] = None,
    paradata_ref: str = "",
    out_dir: str = ".",
    strict: bool = False,
) -> str:
    """A → B → C → D for one file. Returns the written record's path."""
    document = normalize(extract(input_path, doc_id=doc_id))
    record, page_rows, line_rows = to_record(
        document,
        baseline=baseline,
        run_id=run_id,
        paradata_ref=paradata_ref,
        out_dir=out_dir,
        strict=strict,
    )
    return emit(record, page_rows, line_rows, out_path=out_path)


# ──────────────────────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────────────────────


def build_parser() -> argparse.ArgumentParser:
    """Exposed separately so `--help` is testable in-process (repo convention I1)."""
    parser = argparse.ArgumentParser(
        prog="digital_to_json.py",
        description="Convert a digital-born PDF/DOCX into an atrium_document JSON record.",
    )
    parser.add_argument("input", help="path to the .pdf or .docx to convert")
    # W9: `--document-json-out` is the CANONICAL name, `--out` a retained alias.
    #
    # Every other stage in the ecosystem takes the PAIR --document-json /
    # --document-json-out (alto-postprocess, translator, page-classification,
    # nlp-enrich, llm-enrich). This converter took --document-json in and `--out`
    # out, so the one asymmetric spelling in the ecosystem sat on the newest tool —
    # and a pipeline step written by pattern-matching the others would silently write
    # to the default path instead of the requested one.
    #
    # `--out` still works: it is in digital_born/README.md and in the existing tests.
    # argparse resolves both to `args.out` via the shared dest.
    parser.add_argument(
        "--document-json-out",
        "--out",
        dest="out",
        default=None,
        help="output path for the record (default <out-dir>/<doc_id>.document.json). "
        "`--out` is a retained alias for the canonical --document-json-out.",
    )
    parser.add_argument(
        "--document-json",
        default=None,
        help="previous version of the record to accrete onto (rule 1); missing is not an error",
    )
    parser.add_argument("--doc-id", default=None, help="override the derived doc_id")
    parser.add_argument("--out-dir", default=".", help="directory for the default output path")
    parser.add_argument("--run-id", default=None, help="paradata run id to stamp blocks with")
    parser.add_argument("--paradata-ref", default="", help="paradata record this run wrote")
    parser.add_argument(
        "--strict",
        action="store_true",
        help="raise instead of warning on an ownership violation",
    )
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        written = convert(
            args.input,
            out_path=args.out,
            baseline=args.document_json,
            doc_id=args.doc_id,
            run_id=args.run_id,
            paradata_ref=args.paradata_ref,
            out_dir=args.out_dir,
            strict=args.strict,
        )
    except RuntimeError as exc:  # a missing optional dependency, reported as advice
        print(f"[digital-convert] {exc}", file=sys.stderr)
        return 2
    print(f"[digital-convert] wrote {written}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
