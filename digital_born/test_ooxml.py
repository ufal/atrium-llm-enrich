from docx import Document
from docx2python import docx2python


def explore_python_docx(docx_path):
    print(f"--- Running OOXML Exploration on {docx_path} ---")

    # 1. Using python-docx for styles and paragraph metadata
    print("\n[python-docx] Paragraph Styles:")
    doc = Document(docx_path)
    for p_idx, para in enumerate(doc.paragraphs[:5]):  # Inspect first 5 paras
        print(f"Para {p_idx} | Style: {para.style.name} | Text: '{para.text[:50]}...'")

    # Check headers in the first section
    if doc.sections:
        header = doc.sections[0].header
        print(
            f"\n[python-docx] First Section Header Text: '{header.paragraphs[0].text if header.paragraphs else 'None'}'")

    # 2. Using docx2python for a flatter, raw extraction (often catches things python-docx misses)
    print("\n[docx2python] Raw Structure:")
    doc_raw = docx2python(docx_path)
    # docx2python returns a tuple structure: (header, body, footer, footnotes, endnotes)

    if doc_raw.header:
        print(f"Header content found: {doc_raw.header[0][0][0][:50]}...")

    # Print document core properties
    print(f"\nCore Properties: {doc_raw.core_properties}")


if __name__ == "__main__":
    explore_python_docx("sample.docx")
